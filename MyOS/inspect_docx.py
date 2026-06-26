from pathlib import Path
from docx import Document


doc_path = Path("Chapter3_redesign_source.docx")
doc = Document(doc_path)

for i, p in enumerate(doc.paragraphs):
    text = " ".join(p.text.split())
    if not text:
        continue
    style = p.style.name if p.style is not None else ""
    marker = ""
    low = text.lower()
    if "chương 3" in low or "chuong 3" in low or "chương 4" in low or "chuong 4" in low:
        marker = "  <===="
    if marker or style.lower().startswith("heading") or "chương" in low or "chuong" in low:
        print(f"{i:04d} | {style:24} | {text[:220]}{marker}")

print("\n--- Around Chapter 3 ---")
indices = [
    i for i, p in enumerate(doc.paragraphs)
    if "chương 3" in p.text.lower() or "chuong 3" in p.text.lower()
]
if indices:
    start = max(0, indices[0] - 10)
    end = min(len(doc.paragraphs), indices[0] + 180)
    for i in range(start, end):
        text = " ".join(doc.paragraphs[i].text.split())
        if text:
            print(f"{i:04d} | {doc.paragraphs[i].style.name:24} | {text[:260]}")
else:
    print("No Chapter 3 marker found.")
