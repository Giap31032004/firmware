from __future__ import annotations

import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def paragraph_text(p) -> str:
    return norm(p.text)


def read_xml(zf: zipfile.ZipFile, name: str):
    try:
        return ET.fromstring(zf.read(name))
    except KeyError:
        return None


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("usage: analyze_docx_report.py input.docx output.json")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    doc = Document(input_path)

    paragraphs = []
    heading_counts = Counter()
    for idx, p in enumerate(doc.paragraphs):
        text = paragraph_text(p)
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        if style.startswith("Heading") or re.match(r"^(CHƯƠNG|Chương|CHUONG|Chuong|\d+(\.\d+)*)", text):
            heading_counts[style or "plain-numbered"] += 1
        paragraphs.append({"index": idx, "style": style, "text": text})

    table_summaries = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows[:5]:
            rows.append([norm(cell.text) for cell in row.cells])
        table_summaries.append(
            {
                "index": ti,
                "row_count": len(table.rows),
                "column_count": len(table.columns),
                "sample": rows,
            }
        )

    keywords = {
        "workflow": ["workflow", "luồng", "quy trình", "sơ đồ", "flow"],
        "figures": ["hình", "figure", "minh họa", "sơ đồ"],
        "code": ["code", "mã nguồn", "hàm ", "file ", ".c", ".h", "struct", "typedef"],
        "tests": ["kiểm thử", "test", "thử nghiệm", "kết quả", "renode", "uart"],
        "requirements": ["yêu cầu", "mục tiêu", "phạm vi"],
        "limitations": ["hạn chế", "chưa", "tồn tại", "khó khăn"],
        "references": ["tài liệu tham khảo", "tham khảo", "reference"],
    }
    keyword_hits = {k: [] for k in keywords}
    for p in paragraphs:
        low = p["text"].lower()
        for group, words in keywords.items():
            if any(word in low for word in words):
                keyword_hits[group].append(p)

    with zipfile.ZipFile(input_path) as zf:
        names = zf.namelist()
        media = [n for n in names if n.startswith("word/media/")]
        comments_root = read_xml(zf, "word/comments.xml")
        comments = []
        if comments_root is not None:
            for c in comments_root.findall(".//w:comment", NS):
                text = norm("".join(t.text or "" for t in c.findall(".//w:t", NS)))
                comments.append({"author": c.attrib.get(f"{{{NS['w']}}}author", ""), "text": text})

        doc_root = read_xml(zf, "word/document.xml")
        drawing_count = 0
        alt_texts = []
        toc_field_count = 0
        section_count = 0
        if doc_root is not None:
            drawing_count = len(doc_root.findall(".//w:drawing", NS))
            section_count = len(doc_root.findall(".//w:sectPr", NS))
            toc_field_count = sum(
                1
                for instr in doc_root.findall(".//w:instrText", NS)
                if instr.text and "TOC" in instr.text.upper()
            )
            for docpr in doc_root.findall(".//wp:docPr", NS):
                alt_texts.append(
                    {
                        "name": docpr.attrib.get("name", ""),
                        "descr": docpr.attrib.get("descr", ""),
                        "title": docpr.attrib.get("title", ""),
                    }
                )

    text_all = "\n".join(p["text"] for p in paragraphs)
    result = {
        "file": str(input_path),
        "paragraph_count": len(paragraphs),
        "word_count_estimate": len(re.findall(r"\w+", text_all, flags=re.UNICODE)),
        "heading_counts": dict(heading_counts),
        "first_80_paragraphs": paragraphs[:80],
        "all_headings": [
            p
            for p in paragraphs
            if p["style"].startswith("Heading")
            or re.match(r"^(CHƯƠNG|Chương|CHUONG|Chuong|\d+(\.\d+)*)", p["text"])
        ],
        "tables": table_summaries,
        "media_count": len(media),
        "media_files": media,
        "drawing_count": drawing_count,
        "image_alt_texts": alt_texts,
        "toc_field_count": toc_field_count,
        "section_count": section_count,
        "comments": comments,
        "keyword_hit_counts": {k: len(v) for k, v in keyword_hits.items()},
        "keyword_samples": {k: v[:12] for k, v in keyword_hits.items()},
    }

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
